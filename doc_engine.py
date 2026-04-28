from io import BytesIO

import streamlit as st
from docx import Document
from docx.shared import Inches


def create_document(title=None):
    try:
        document = Document()
        if title:
            document.add_heading(str(title), level=1)
        return document
    except Exception as exc:
        st.error(f"Error in doc_engine.py at function create_document: {exc}")
        return Document()


def add_paragraph(document, text, style=None):
    try:
        return document.add_paragraph(str(text or ""), style=style)
    except Exception as exc:
        st.error(f"Error in doc_engine.py at function add_paragraph: {exc}")
        return None


def add_heading(document, text, level=2):
    try:
        return document.add_heading(str(text or ""), level=level)
    except Exception as exc:
        st.error(f"Error in doc_engine.py at function add_heading: {exc}")
        return None


def add_table(document, headers, rows):
    try:
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = str(header)

        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row[:len(headers)]):
                cells[index].text = str(value)
        return table
    except Exception as exc:
        st.error(f"Error in doc_engine.py at function add_table: {exc}")
        return None


def add_image(document, image_bytes, width_inches=5.5):
    try:
        image_stream = BytesIO(image_bytes)
        return document.add_picture(image_stream, width=Inches(width_inches))
    except Exception as exc:
        st.error(f"Error in doc_engine.py at function add_image: {exc}")
        return None


def save_document_bytes(document):
    try:
        output = BytesIO()
        document.save(output)
        return output.getvalue()
    except Exception as exc:
        st.error(f"Error in doc_engine.py at function save_document_bytes: {exc}")
        return b""


def generate_summary_document(title, sections=None, tables=None, images=None):
    try:
        document = create_document(title)

        for heading, body in (sections or []):
            add_heading(document, heading, level=2)
            add_paragraph(document, body)

        for table_config in (tables or []):
            add_heading(document, table_config.get("title", "Table"), level=2)
            add_table(document, table_config.get("headers", []), table_config.get("rows", []))

        for image_config in (images or []):
            add_heading(document, image_config.get("title", "Image"), level=2)
            add_image(document, image_config.get("bytes", b""), image_config.get("width_inches", 5.5))

        return save_document_bytes(document)
    except Exception as exc:
        st.error(f"Error in doc_engine.py at function generate_summary_document: {exc}")
        return b""
